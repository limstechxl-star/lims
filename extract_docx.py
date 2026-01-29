from docx import Document

doc = Document('Labax.docx')

# Write to file with UTF-8 encoding
with open('docx_content.txt', 'w', encoding='utf-8') as f:
    f.write("=" * 80 + "\n")
    f.write("DOCUMENT CONTENT ANALYSIS\n")
    f.write("=" * 80 + "\n")

    f.write("\n--- PARAGRAPHS ---\n\n")
    for i, para in enumerate(doc.paragraphs, 1):
        if para.text.strip():
            f.write(f"[{i}] {para.text}\n")

    f.write("\n--- TABLES ---\n")
    for i, table in enumerate(doc.tables, 1):
        f.write(f"\nTable {i}:\n")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            f.write(" | ".join(row_data) + "\n")

    f.write("\n--- DOCUMENT STATISTICS ---\n\n")
    f.write(f"Total Paragraphs: {len(doc.paragraphs)}\n")
    f.write(f"Total Tables: {len(doc.tables)}\n")
    f.write(f"Total Sections: {len(doc.sections)}\n")

print("Content extracted to docx_content.txt")
